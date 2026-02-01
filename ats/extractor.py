from __future__ import annotations
import os
from typing import Tuple, Dict, Any

from .utils import normalize_spaces

def _extract_docx(path: str) -> Tuple[str, Dict[str, Any]]:
    from docx import Document  # python-docx
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    # include table text (but warn: tables can be ATS-unfriendly)
    table_cells = 0
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                table_cells += 1
                txt = (cell.text or "").strip()
                if txt:
                    parts.append(txt)

    meta = {
        "file_type": "docx",
        "has_tables": len(doc.tables) > 0,
        "table_cells": table_cells,
    }
    return normalize_spaces("\n".join(parts)), meta

def _extract_pdf_pdfplumber(path: str) -> Tuple[str, Dict[str, Any]]:
    import pdfplumber
    texts = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            texts.append(p.extract_text() or "")
    txt = normalize_spaces("\n".join(texts))
    meta = {
        "file_type": "pdf",
        "pages": len(texts),
        "backend": "pdfplumber",
    }
    return txt, meta

def extract_text(path: str) -> Tuple[str, Dict[str, Any]]:
    ext = os.path.splitext(path)[1].lower().strip(".")
    if ext == "docx":
        return _extract_docx(path)
    if ext == "pdf":
        return _extract_pdf_pdfplumber(path)
    raise ValueError("Format non supporté. Utilise .pdf ou .docx")
